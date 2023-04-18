import os, subprocess, shutil, uuid
from docx import Document
from custom_logging import logger

if os.name == 'nt':
    import comtypes.client


    wdFormatPDF = 17


    def covx_to_pdf(infile, outfile):
        """Convert a Word .docx to PDF"""

        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(infile)
        doc.SaveAs(outfile, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()


'''
    list_variables -> [(name, bold, italic), (age, bold italic), (institution, bold, italic)]
    valores -> ['Luis Bagner', '22', 'random institution']
'''
def generar_archivo(path_plantilla:str, path_salida:str, formato_inicio:str, formato_fin:str, pdf_or_word:str, list_variables:list, valores:list) -> None:
    try:
        document = Document(path_plantilla)
        
        #name;age;institution
        #Luis Bagner;22;random institution
        for paragraph in document.paragraphs:
            text = paragraph.text
            pre_build = []

            for index in range(len(list_variables)):
                variable = formato_inicio + list_variables[index][0] + formato_fin
                if variable in text:
                    paragraph = paragraph.clear()
                    part1, text = text.split(variable)[0], text.split(variable)[1]
                    pre_build.append(part1)
                    pre_build.append((valores[index], list_variables[index][1], list_variables[index][2]))
            if len(pre_build) > 0:
                if type(pre_build[-1]) == tuple and text != '':
                    pre_build.append(text)
                for line in pre_build:
                    if type(line) == tuple:
                        new_paragraph = paragraph.add_run(line[0])
                        new_paragraph.bold = line[1]
                        new_paragraph.italic = line[2]
                    else:
                        paragraph.add_run(line)

        file_name = uuid.uuid4().__str__()

        if pdf_or_word == 'PDF':
            temp_path = os.path.normpath(os.path.join(path_salida, f'.{file_name}.docx'))
            temp_path_pre_pdf = os.path.normpath(os.path.join(path_salida, f'.{file_name}.pdf'))
            temp_path_to_pdf = os.path.normpath(os.path.join(path_salida, f'{file_name}.pdf'))

            document.save(temp_path)
            # convert pdf to docx
            if os.name == 'posix':
                subprocess.run(['libreoffice', '--convert-to', 'pdf', temp_path, '--outdir', path_salida, '--headless'])
                subprocess.run(['mv', temp_path_pre_pdf, temp_path_to_pdf])
            elif os.name == 'nt':
                covx_to_pdf(temp_path, temp_path_to_pdf)
            os.remove(temp_path)
        elif pdf_or_word == 'Word':
            temp_path = os.path.join(path_salida, f'{file_name}.docx')
            document.save(temp_path)
    except Exception as ex:
        logger.error(ex)


def generar_archivos(path_plantilla:str, path_salida:str, formato_inicio:str, formato_fin:str, pdf_or_word:str, list_variables:list,  list_valores:list) -> None:
    
    # Eliminar directorio completo antes de generar los archivos
    path_salida = os.path.normpath(os.path.join(path_salida, 'generated_files'))
    try:
        shutil.rmtree(path_salida)
    except FileNotFoundError as ex:
        logger.info(ex)
    try:
        os.mkdir(path_salida)   
    except FileExistsError as ex:
        logger.info(ex)

    for valores in list_valores:
        generar_archivo(path_plantilla, path_salida, formato_inicio, formato_fin, pdf_or_word, list_variables, valores)

